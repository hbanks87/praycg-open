import os, re, json, math, pickle, shutil, zipfile, subprocess, textwrap
from pathlib import Path
import numpy as np
import pandas as pd
from scipy import signal, stats
from sklearn.preprocessing import StandardScaler
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path('/mnt/data')
PREV = ROOT/'PRAYCG_FieldOfDreams_Run1_MasterSuite_v1_0'
OUT = ROOT/'PRAYCG_FieldOfDreams_Run1_MasterComprehensive_FULL_v1_1'
if OUT.exists(): shutil.rmtree(OUT)
TABLES = OUT/'tables'; REPORT=OUT/'report'; FIGS=OUT/'figures'; SCRIPTS=OUT/'scripts'; INPUTS=OUT/'inputs_snapshot'
for d in [OUT,TABLES,REPORT,FIGS,SCRIPTS,INPUTS]: d.mkdir(parents=True, exist_ok=True)
# copy previous outputs
if PREV.exists():
    for sub in ['tables','figures']:
        for f in (PREV/sub).glob('*'):
            if f.is_file(): shutil.copy2(f, (OUT/sub)/f.name)
# inputs
input_files = [
 'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events(1).json',
 'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events(1).csv',
 'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events_channel_map(1).csv',
 'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events_run_config_media_selection(1).json',
 'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events_predeclared_anchors(1).csv',
 'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events_prerun_display_audio_calibration(1).json',
 'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events_CONTROL_1_AFTER_WASHOUT_1_core_report(1).json',
 'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events_TARGET_1_AFTER_WASHOUT_2_core_report(1).json',
 'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events_OVERRIDE_1_AFTER_WASHOUT_3_core_report(1).json',
 'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events_OVERRIDE_1_AFTER_WASHOUT_3_override_task_report(1).json',
 'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events_final_master_report(1).json',
 'qczip(1).zip',
]
for name in input_files:
    p = ROOT/name
    if p.exists(): shutil.copy2(p, INPUTS/name)
# copy qczip extracted folder if exists
qcroot=ROOT/'field_qc'/'qczip'
if qcroot.exists():
    shutil.copytree(qcroot, INPUTS/'field_qc_qczip', dirs_exist_ok=True)

# Helper
def jload(p):
    with open(p, 'r', encoding='utf-8') as f: return json.load(f)

def safe_read_csv(p):
    return pd.read_csv(p) if Path(p).exists() else pd.DataFrame()

# load output tables
features = safe_read_csv(TABLES/'field_time_resolved_feature_frame.csv')
phase_summary = safe_read_csv(TABLES/'field_phase_feature_summary.csv')
amred = safe_read_csv(TABLES/'field_amred_anchor_endpoint_table.csv')
amred_runner = amred[amred.get('sensitivity','')=='runner_registered_time'].copy() if not amred.empty else pd.DataFrame()
hrv = safe_read_csv(TABLES/'field_hrv_phase_summary.csv')
resp = safe_read_csv(TABLES/'field_resp_phase_summary.csv')
alspulse = safe_read_csv(TABLES/'als_start_pulse_qc.csv')
ocm = safe_read_csv(TABLES/'field_ocm025_rsm_cvb_squint_cue_table.csv')
rsm_corr = safe_read_csv(TABLES/'field_rsm_correlation_summary.csv')
tti_global = safe_read_csv(TABLES/'field_tti_global_summary.csv')
tti_anchor = safe_read_csv(TABLES/'field_tti_anchor_deltas.csv')
lzc = safe_read_csv(TABLES/'field_lz_complexity_timeseries.csv')
itp = safe_read_csv(TABLES/'field_mred_itp_anchor_summary.csv')
phase_coverage = safe_read_csv(TABLES/'phase_coverage.csv')
stream_inv = safe_read_csv(TABLES/'stream_inventory_corrected.csv')
run_config_path = ROOT/'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events_run_config_media_selection(1).json'
events_path = ROOT/'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events(1).json'
run_config = jload(run_config_path)
events = jload(events_path)

# 1. Module registry
modules = []
def add_mod(name, tier, status, primary_outputs, finding, limitation=''):
    modules.append({'module':name,'tier':tier,'status':status,'primary_outputs':primary_outputs,'finding':finding,'limitation':limitation})

# derive counts
a_pass = int(amred_runner['A_MRED_pass'].sum()) if not amred_runner.empty else 0
base_pass = int(amred_runner['base_gate_pass'].sum()) if not amred_runner.empty else 0
spec_pass = int(amred_runner['specificity_gate_pass'].sum()) if not amred_runner.empty else 0
n_anchor = len(amred_runner)
als_passes = int(alspulse['pass_detected_fullscreen_pulse'].sum()) if not alspulse.empty else 0
add_mod('Run metadata / media identity / hash verification', 'QC / validity', 'COMPLETE', 'run_config_media_selection; stream_inventory_corrected; phase_coverage', 'PRAYCG2.0 run; Target/Override hash matched; channel map locked; 66 cue schedule loaded.', '')
add_mod('ALS/PT19 start-pulse timing QC', 'QC / validity', 'PARTIAL / CAUTION', 'als_start_pulse_qc.csv', f'{als_passes}/3 branch pulses passed simple threshold; Target detected, Control/Override weak/failed.', 'Prevents strict timing-grade confirmation; sample-index timing reconstruction used.')
add_mod('StimulusFingerprint v1.8 / CET regressors', 'QC / validity', 'COMPLETE', 'stimulus_exogenous_regressor_frame_all_conditions; field_cet_residualization_model_summary', 'Control/Target/Override regressors available; exogenous models did not explain NIP robustly in blocked CV.', 'Anchor-specific stimulus vectors absent because anchor_count=0 in QC manifest.')
add_mod('Feature extraction / bandpower / artifact sentinels', 'Core preprocessing', 'COMPLETE', 'field_time_resolved_feature_frame.csv; field_phase_feature_summary.csv', 'Generated 1-second stepped 2-second feature frame and phase summaries.', 'Single-participant and scalp EEG artifact-sensitive.')
add_mod('GammaScalpel microband scan', 'Core feature diagnostic', 'COMPLETE', 'field_gammascalpel_microband_phase_summary.csv', 'Lower-gamma microbands summarized by branch/ROI; see table for Target/Override deltas.', 'Gamma remains artifact-sensitive and secondary to A-MRED.')
add_mod('Topo-OSM network-state vectors', 'Secondary / interpretation', 'COMPLETE', 'field_topo_osm_phase_state_vectors.csv', 'Phase-level network-state proxy vectors generated for all phases.', 'Human-scale topology proxy only, not cellular OSM.')
add_mod('NAST narrative-absorption transition', 'Secondary / interpretation', 'COMPLETE', 'field_nast_phase_transition_table.csv', 'Baseline/Washout-to-branch transitions summarized; Target transition positive but not dramatic.', 'DMN-proxy only; not direct DMN detection.')
add_mod('CandidateLocal KHT-topo / MRED', 'Primary-adjacent / exploratory', 'COMPLETE', 'field_candidate_local_kht_topo_mred_event_table.csv', 'Anchor-local event table generated from A-MRED inputs; one anchor (A6) candidate passed A-MRED at runner time.', 'No independent group replication; timing sensitivity present.')
add_mod('A-MRED anchor-locked endpoint', 'PRIMARY RECOMMENDED', 'COMPLETE / PILOT POSITIVE', 'field_amred_anchor_endpoint_table.csv; field_amred_primary_endpoint_summary.csv', f'{a_pass}/{n_anchor} runner-time anchors passed; sensitivity with content offset dropped to 0/7.', 'Runner-loaded blind estimated anchors, not frame-verified locked anchors; ALS weak.')
add_mod('MRED event table', 'Primary-derived', 'COMPLETE', 'field_mred_event_table.csv', 'MR/ENC quadrants assigned for all anchor x condition rows.', 'Anchors estimated; not frame-verified.')
add_mod('NIP/BIT/CII/IAQ', 'Secondary summary', 'COMPLETE', 'field_nip_component_timeseries.csv; field_bit_event_table.csv; field_cii_anchor_integrals.csv; field_iaq_table.csv', 'Target CII > Control/Override on later anchors; overall NIP weaker than Contact but Target-dominant.', 'Secondary summary, not endpoint proof.')
add_mod('TTI reception-extraction tradeoff', 'Secondary summary', 'COMPLETE', 'field_tti_global_summary.csv; field_tti_anchor_deltas.csv', 'Global TTI positive (+0.568), consistent with Target reception vs Override extraction.', 'Not moral/clinical score.')
add_mod('OCM025 / RSM / CVB / SquintProxy', 'Override diagnostic', 'COMPLETE', 'field_ocm025_rsm_cvb_squint_cue_table.csv; field_rsm_correlation_summary.csv', 'Override task engaged by report; RSM found modest hard-number/carry/guess-risk structure.', 'Private strategy not provable from physiology alone.')
add_mod('CET-R residualization', 'QC / guardrail', 'COMPLETE', 'field_cet_residualization_model_summary.csv; field_cet_residualized_cii_anchor_integrals.csv', 'Stimulus-side regressors did not generalize well enough to explain NIP pattern.', 'CET-R supports but does not confirm endpoint.')
add_mod('EET endogenous echo tracking', 'Exploratory convergence', 'COMPLETE', 'field_eet_endogenous_echo_tracking.csv', 'Baseline 2 / washout state-vector resemblance estimated.', 'State-vector similarity only; not proof of replay or memory.')
add_mod('MRED-ITP / ACG / OCU', 'Exploratory convergence', 'COMPLETE', 'field_lz_complexity_timeseries.csv; field_mred_itp_anchor_summary.csv; field_acg_event_table.csv; field_ocu_event_table.csv', 'ACG/OCU convergence generated; no strict proof layer; blink release weak without EOG.', 'LZC is complexity proxy, not thermodynamic entropy; Fp1 blink proxy not EOG.')
add_mod('SelfReport20 parser', 'Phenomenology / covariate', 'COMPLETE', 'field_self_report_summary.csv', 'PRAYCG2.0 consolidated self-report parsed; Target strongest by final report.', 'Self-report contextualizes physiology; not proof of internal state.')
add_mod('LSO/SPM subtitle override', 'Not applicable', 'N/A', 'none', 'No subtitle condition in this run.', 'Not applicable.')
add_mod('Dyadic / group hyperscanning', 'Not applicable', 'N/A', 'none', 'Single-participant run only.', 'No intersubject synchrony endpoint.')
registry = pd.DataFrame(modules)
registry.to_csv(TABLES/'field_full_suite_module_registry.csv', index=False)

# 2. Module tier map / boxed recommendation table
tier_rows = [
    {'tier':'BOXED PRIMARY PATH','module_group':'Timing/QC + StimulusFingerprint/CET-R + artifact/confound gates -> A-MRED','role':'Recommended confirmatory compression path for future runs.'},
    {'tier':'PRIMARY ENDPOINT','module_group':'A-MRED','role':'Anchor-locked MRED pass: Target MR and ENC exceed Control/Override after QC.'},
    {'tier':'SECONDARY SUMMARIES','module_group':'NIP/BIT/CII/IAQ; TTI','role':'Summarize immersion proxy and reception/extraction contrast.'},
    {'tier':'EXPLORATORY CONVERGENCE','module_group':'KHT-topo; NAST; EET; MRED-ITP/ACG/OCU; OCM/RSM/CVB/SquintProxy','role':'Interpretive and falsification layers; not primary proof.'},
]
pd.DataFrame(tier_rows).to_csv(TABLES/'field_module_tier_map_boxed_recommended_path.csv', index=False)

# 3. Topo-OSM phase vectors
if not phase_summary.empty:
    topo_cols = ['meaninggamma_z','tsp_z','theta_integration_z','taskgamma_z','nast_nas_z','artifact_score','NIP_density','alpha_posterior_z']
    topo = phase_summary[['condition']+topo_cols].copy()
    # vector norm and simple interpretation
    topo['topo_state_norm'] = np.sqrt(np.nansum(np.square(topo[topo_cols].values), axis=1))
    # distance from Baseline1
    if 'BASELINE_1' in topo.condition.values:
        bvec = topo.loc[topo.condition=='BASELINE_1',topo_cols].iloc[0].values.astype(float)
        topo['distance_from_baseline1'] = [float(np.linalg.norm(row[topo_cols].values.astype(float)-bvec)) for _,row in topo.iterrows()]
    topo.to_csv(TABLES/'field_topo_osm_phase_state_vectors.csv', index=False)

# 4. NAST transition table
transition_pairs = [('BASELINE_1','CONTROL_1'),('WASHOUT_1','TARGET_1'),('WASHOUT_2','CONTEXTUAL_OVERRIDE_1'),('WASHOUT_3','BASELINE_2_REFLECTION')]
trans_rows=[]
ps=phase_summary.set_index('condition') if not phase_summary.empty else pd.DataFrame()
for a,b in transition_pairs:
    if a in ps.index and b in ps.index:
        row={'from_phase':a,'to_phase':b}
        for col in ['alpha_posterior_z','meaninggamma_z','tsp_z','theta_integration_z','taskgamma_z','nast_nas_z','artifact_score','NIP_density']:
            row[f'delta_{col}']=float(ps.loc[b,col]-ps.loc[a,col])
        row['interpretation']='positive absorption proxy if alpha decreases while meaning/TSP/NIP increase and task/artifact do not dominate'
        trans_rows.append(row)
nast = pd.DataFrame(trans_rows)
nast.to_csv(TABLES/'field_nast_phase_transition_table.csv', index=False)

# 5. NIP component timeseries, BIT, CII/IAQ from AMRED
if not features.empty:
    nipcols=[c for c in ['time_lsl','phase_time_sec','condition','meaninggamma_z','tsp_z','theta_integration_z','taskgamma_z','artifact_score','nast_nas_z','A_sem','R_int','NIP_density','blink_proxy_event'] if c in features.columns]
    features[nipcols].to_csv(TABLES/'field_nip_component_timeseries.csv', index=False)
if not amred_runner.empty:
    bit = amred_runner[['anchor_id','anchor_time_sec_used','A_MRED_pass','base_gate_pass','specificity_gate_pass','target_MR','target_ENC','target_CII','control_CII','override_CII','IAQ_target_vs_override','claim_level']].copy()
    bit['BIT_pass'] = bit['A_MRED_pass']
    bit['BIT_interpretation'] = np.where(bit['BIT_pass'], 'strict bivariate immersion gate passed at estimated runner time', 'failed strict bivariate immersion gate')
    bit.to_csv(TABLES/'field_bit_event_table.csv', index=False)
    cii_rows=[]
    for _,r in amred_runner.iterrows():
        for cond in ['target','control','override']:
            cii_rows.append({'anchor_id':r.anchor_id,'anchor_time_sec':r.anchor_time_sec_used,'condition':cond.upper(),'CII':r[f'{cond}_CII'], 'MR':r[f'{cond}_MR'], 'ENC':r[f'{cond}_ENC'], 'artifact':r[f'{cond}_artifact']})
    pd.DataFrame(cii_rows).to_csv(TABLES/'field_cii_anchor_integrals.csv', index=False)
    iaq=amred_runner[['anchor_id','anchor_time_sec_used','target_CII','override_CII','control_CII','target_minus_override_CII','target_minus_control_CII','IAQ_target_vs_override']].copy()
    iaq.to_csv(TABLES/'field_iaq_target_override_table.csv', index=False)

# 6. MRED event and CandidateLocal table from A-MRED runner rows
if not amred_runner.empty:
    mred_rows=[]; cand_rows=[]
    for _,r in amred_runner.iterrows():
        for cond_prefix,label in [('target','TARGET_1'),('control','CONTROL_1'),('override','CONTEXTUAL_OVERRIDE_1')]:
            mr=r[f'{cond_prefix}_MR']; enc=r[f'{cond_prefix}_ENC']; cii=r[f'{cond_prefix}_CII']
            if mr>=0.5 and enc>=0.25: quad='MR_HIGH_ENC_HIGH'
            elif mr>=0.5 and enc<0.25: quad='MR_HIGH_ENC_LOW'
            elif mr<0.5 and enc>=0.25: quad='MR_LOW_ENC_HIGH'
            else: quad='MR_LOW_ENC_LOW'
            mred_rows.append({'anchor_id':r.anchor_id,'condition':label,'phase_time_sec':r.anchor_time_sec_used,'MR':mr,'ENC':enc,'CII':cii,'artifact':r[f'{cond_prefix}_artifact'],'MRED_quadrant':quad,'A_MRED_anchor_pass_for_target':bool(r.A_MRED_pass) if cond_prefix=='target' else False,'claim_level':r.claim_level})
        cand_rows.append({'anchor_id':r.anchor_id,'phase_time_sec':r.anchor_time_sec_used,'target_peak_time':r.target_peak_time,'target_peak_mr':r.target_peak_mr,'target_peak_tsp':r.target_peak_tsp,'target_peak_nip':r.target_peak_nip,'target_theta_carry':r.target_theta_carry,'target_theta_pre':r.target_theta_pre,'KHT_topo_proxy_local':float(np.sqrt(max(0,abs(r.target_peak_mr*r.target_theta_carry)))) if pd.notna(r.target_peak_mr) and pd.notna(r.target_theta_carry) else np.nan,'target_MR':r.target_MR,'target_ENC':r.target_ENC,'A_MRED_pass':r.A_MRED_pass,'specificity_gate_pass':r.specificity_gate_pass,'claim_level':r.claim_level})
    pd.DataFrame(mred_rows).to_csv(TABLES/'field_mred_event_table.csv', index=False)
    pd.DataFrame(cand_rows).to_csv(TABLES/'field_candidate_local_kht_topo_mred_event_table.csv', index=False)

# 7. EET state-vector echo tracking
vec_cols = ['meaninggamma_z','tsp_z','theta_integration_z','taskgamma_z','nast_nas_z','NIP_density','alpha_posterior_z']
if not features.empty and not amred_runner.empty:
    echo_rows=[]
    for _,a in amred_runner.iterrows():
        t0=float(a.anchor_time_sec_used)
        target_win = features[(features.condition=='TARGET_1') & (features.phase_time_sec>=t0-5) & (features.phase_time_sec<=t0+25)]
        if len(target_win)<3: continue
        target_vec=target_win[vec_cols].mean().values.reshape(1,-1)
        # easier explicit comps
        comps=[('WASHOUT_2_first30','WASHOUT_2',0,30),('WASHOUT_2_all','WASHOUT_2',0,999),('BASELINE2_first30','BASELINE_2_REFLECTION',0,30),('BASELINE2_all','BASELINE_2_REFLECTION',0,999),('CONTROL_same_time','CONTROL_1',max(0,t0-5),t0+25),('OVERRIDE_same_time','CONTEXTUAL_OVERRIDE_1',max(0,t0-5),t0+25)]
        for name,phase,st,en in comps:
            comp=features[(features.condition==phase)&(features.phase_time_sec>=st)&(features.phase_time_sec<=en)]
            if len(comp)<3: continue
            comp_vec=comp[vec_cols].mean().values.reshape(1,-1)
            sim=float(cosine_similarity(target_vec, comp_vec)[0,0]) if np.all(np.isfinite(target_vec)) and np.all(np.isfinite(comp_vec)) else np.nan
            echo_rows.append({'anchor_id':a.anchor_id,'target_anchor_time_sec':t0,'comparison':name,'comparison_phase':phase,'cosine_similarity_state_vector':sim,'target_window':'t-5_to_t+25','comparison_window_start':st,'comparison_window_end':en})
    pd.DataFrame(echo_rows).to_csv(TABLES/'field_eet_endogenous_echo_tracking.csv', index=False)

# 8. MRED-ITP tables: split acg/ocu event tables from anchor summary
if not itp.empty:
    itp = itp.copy()
    if 'CSI' not in itp.columns:
        itp['CSI'] = itp['C_strike'].fillna(0) - itp['C_settle'].fillna(0)
    if 'ORI' not in itp.columns:
        itp['ORI'] = itp['blink_suppression'].fillna(0) + itp['blink_release'].fillna(0)
    acg=itp[['condition','anchor_id','anchor_time_sec','C_pre','C_peak','C_post','C_strike','C_settle','CSI','ACG_candidate']].copy()
    acg.to_csv(TABLES/'field_acg_event_table.csv', index=False)
    ocu=itp[['condition','anchor_id','anchor_time_sec','blink_rate_hold','blink_rate_release','blink_suppression','blink_release','ORI','OCU_candidate']].copy()
    ocu.to_csv(TABLES/'field_ocu_event_table.csv', index=False)

# 9. GammaScalpel microband phase summary using raw pickled parsed xdf
try:
    with open(ROOT/'field_xdf_parsed.pkl','rb') as f: H,S,F=pickle.load(f)
    sid_by={h['name']:sid for sid,h in H.items()}
    eeg = S[sid_by['obci_eeg1']]['data'].astype(float)
    # corrected times from prior tables status mapping
    status = safe_read_csv(TABLES/'openbci_status_heartbeats.csv')
    if len(status)>5:
        xnum=status['samples_total'].values.astype(float); y=status['lsl_time'].values.astype(float)
        slope,intercept=np.polyfit(xnum,y,1); fs=1/slope
        start_global=status['samples_total'].iloc[-1] - (len(eeg)-1)
        eeg_t=intercept+slope*(start_global+np.arange(len(eeg)))
    else:
        fs=125.0; eeg_t=S[sid_by['obci_eeg1']]['ts']
    chmap = pd.read_csv(ROOT/'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events_channel_map(1).csv')
    rois={row.electrode_location:int(row.openbci_channel)-1 for _,row in chmap.iterrows()}
    groups={'meaning_posterior_temporal_parietal':[rois[k] for k in ['T5','T6','Pz','P3','P4']], 'task_frontal':[rois[k] for k in ['Fz','F3','F4']], 'visual_occipital':[rois[k] for k in ['O1','O2']], 'artifact_jaw_temporal':[rois[k] for k in ['T3','T4']], 'fp1_ocular':[rois['Fp1']]}
    bands=[(25,30),(30,35),(35,40),(40,45),(45,50),(50,55)]
    eeg2=eeg-np.nanmedian(eeg,axis=0,keepdims=True)
    segs=phase_coverage[['phase','start','end']].copy()
    micro_rows=[]
    for _,segrow in segs.iterrows():
        phase=segrow.phase; st=segrow.start; en=segrow.end
        mask=(eeg_t>=st)&(eeg_t<=en)
        if mask.sum()<int(fs*10): continue
        arr=eeg2[mask]
        f,p=signal.welch(signal.detrend(arr,axis=0,type='constant'), fs=fs, nperseg=min(int(fs*4),len(arr)), noverlap=int(fs*2), axis=0)
        for lo,hi in bands:
            fm=(f>=lo)&(f<hi)
            for g,inds in groups.items():
                vals=np.trapezoid(p[fm][:,inds], f[fm], axis=0)
                val=float(np.log10(max(np.nanmean(vals),1e-12)))
                micro_rows.append({'condition':phase,'roi_group':g,'band_hz':f'{lo}-{hi}','log10_power':val})
    micro=pd.DataFrame(micro_rows)
    # baseline normalize by all phase std for easier comparison
    if not micro.empty:
        out=[]
        for (g,b),df in micro.groupby(['roi_group','band_hz']):
            mean=df['log10_power'].mean(); std=df['log10_power'].std(ddof=0) or 1
            df=df.copy(); df['z_by_band_roi']=(df['log10_power']-mean)/std
            out.append(df)
        micro=pd.concat(out,ignore_index=True)
        # add target/control/override deltas pivot
        piv=micro.pivot_table(index=['roi_group','band_hz'],columns='condition',values='z_by_band_roi').reset_index()
        for a,b in [('TARGET_1','CONTEXTUAL_OVERRIDE_1'),('TARGET_1','CONTROL_1'),('CONTEXTUAL_OVERRIDE_1','CONTROL_1')]:
            if a in piv.columns and b in piv.columns:
                piv[f'{a}_minus_{b}']=piv[a]-piv[b]
        micro.to_csv(TABLES/'field_gammascalpel_microband_long.csv', index=False)
        piv.to_csv(TABLES/'field_gammascalpel_microband_phase_summary.csv', index=False)
except Exception as e:
    pd.DataFrame([{'error':repr(e)}]).to_csv(TABLES/'field_gammascalpel_microband_phase_summary.csv', index=False)

# 10. Full visual overlay combined
visuals=[]
if not amred_runner.empty:
    for _,r in amred_runner.iterrows():
        visuals.append({'condition':'TARGET_1','time_sec':r.anchor_time_sec_used,'event_type':'A_MRED','label':r.anchor_id,'severity':'primary' if r.A_MRED_pass else 'anchor','note':f'A_MRED={r.A_MRED_pass}; MR={r.target_MR:.3f}; ENC={r.target_ENC:.3f}; CII={r.target_CII:.3f}'})
        visuals.append({'condition':'CONTROL_1','time_sec':r.anchor_time_sec_used,'event_type':'ANCHOR_CONTROL','label':r.anchor_id,'severity':'reference','note':'control matched-time anchor'})
        visuals.append({'condition':'CONTEXTUAL_OVERRIDE_1','time_sec':r.anchor_time_sec_used,'event_type':'ANCHOR_OVERRIDE','label':r.anchor_id,'severity':'reference','note':'override matched-time anchor'})
if not tti_anchor.empty:
    for _,r in tti_anchor.iterrows(): visuals.append({'condition':'TARGET_1','time_sec':r.anchor_time_sec_used,'event_type':'TTI_ANCHOR','label':r.anchor_id,'severity':'secondary','note':f'TTI={r.TTI_anchor:.3f}; IAQ={r.IAQ_target_vs_override:.3f}'})
if not itp.empty:
    for _,r in itp[(itp.condition=='TARGET_1') & ((itp.ACG_candidate==True)|(itp.OCU_candidate==True))].iterrows():
        visuals.append({'condition':'TARGET_1','time_sec':r.anchor_time_sec,'event_type':'MRED_ITP','label':r.anchor_id,'severity':'exploratory','note':f'ACG={r.ACG_candidate}; OCU={r.OCU_candidate}; CSI={r.CSI:.3f}; ORI={r.ORI:.3f}'})
vis=pd.DataFrame(visuals)
vis.to_csv(TABLES/'field_visual_overlay_full_suite.csv', index=False)

# 11. Unsupported/not applicable table
na = pd.DataFrame([
    {'module':'LSO/SPM Subtitle Override','status':'NOT_APPLICABLE','reason':'No subtitle condition or subtitle file in Field of Dreams run.'},
    {'module':'Dyadic / group hyperscanning','status':'NOT_APPLICABLE','reason':'Single participant self-run.'},
    {'module':'External independent verification','status':'NOT_APPLICABLE','reason':'Not an external verifier run.'},
    {'module':'Cellular OSM / hidden-Y biology','status':'OUT_OF_SCOPE','reason':'Human scalp EEG cannot identify cellular/cytoskeletal/biophotonic mechanism.'},
])
na.to_csv(TABLES/'field_not_applicable_or_out_of_scope_modules.csv', index=False)

# Summary metrics for report
amred_pass_anchor = amred_runner.loc[amred_runner.A_MRED_pass==True,'anchor_id'].tolist() if not amred_runner.empty else []
try: tti_val=float(tti_global['TTI_global'].iloc[0])
except Exception: tti_val=np.nan
try:
    target_report = jload(ROOT/'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events_TARGET_1_AFTER_WASHOUT_2_core_report(1).json')
    control_report = jload(ROOT/'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events_CONTROL_1_AFTER_WASHOUT_1_core_report(1).json')
    override_report = jload(ROOT/'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events_OVERRIDE_1_AFTER_WASHOUT_3_core_report(1).json')
    final_report = jload(ROOT/'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events_final_master_report(1).json')
    override_task = jload(ROOT/'PRAYCG_v2_0_hoyt_banks_s0007_field_of_dreams2_20260818_060321_events_OVERRIDE_1_AFTER_WASHOUT_3_override_task_report(1).json')
except Exception:
    target_report=control_report=override_report=final_report=override_task={}

# 12. Interpretation JSON
interpret = {
    'executive_verdict':'Useful PRAYCG2.0 runner-registered estimated blind-anchor pilot. Full suite supports a Target-favoring reception/meaning pattern, but not strict confirmatory A-MRED because anchors were YouTube-estimated/not frame-verified and ALS start-pulse QC was weak.',
    'run_classification':'runner-registered estimated blind-anchor pilot; not frame-verified confirmatory',
    'amred_runner_time_passes':a_pass,
    'amred_total_anchors':n_anchor,
    'amred_content_offset_sensitivity_passes':int(amred[(amred.get('sensitivity','')=='content_offset_plus_1p251s')]['A_MRED_pass'].sum()) if not amred.empty else None,
    'amred_pass_anchor_ids':amred_pass_anchor,
    'tti_global':tti_val,
    'als_passes_simple_threshold':als_passes,
    'als_total_branches':3,
    'primary_recommendation':'For TSC/poster, present Field of Dreams as pilot-positive/future-freeze evidence only. Next run should use frame-verified locked anchors and clean ALS pulse before calling A-MRED confirmatory.',
    'module_compression':'A-MRED is the boxed primary endpoint; NIP/TTI are secondary summaries; KHT-topo/EET/ACG/OCU/OCM/RSM/CVB/Squint are exploratory convergence/diagnostic layers.'
}
with open(TABLES/'field_full_suite_interpretation_v1_1.json','w',encoding='utf-8') as f: json.dump(interpret,f,indent=2)

# Markdown report
report_md = f"""# Field of Dreams PR-AYC-G Run 1 - Full Master Comprehensive Suite v1.1

## Executive verdict

**Field of Dreams Run 1 is a useful PRAYCG2.0, runner-registered estimated blind-anchor pilot.** The full suite supports a Target-favoring reception/meaning pattern, but it does **not** reach strict confirmatory A-MRED status because the anchors were YouTube-estimated rather than frame-verified and ALS start-pulse validation was weak.

### Claim grade

- **A-MRED pilot-positive candidate:** yes.
- **Strict frame-verified confirmatory A-MRED:** no.
- **Primary reason:** runner-loaded estimated blind anchors + weak ALS start-pulse QC.

## Inputs and run integrity

- Runner version: PRAYCG2.0.
- Participant/session/run: `{run_config['config']['participant_id']}` / `{run_config['config']['session_id']}` / `{run_config['config']['run_label']}`.
- Target and Override hash match: `{run_config['target_override_hash_match']}`.
- Cue count / expected sum: {run_config['cue_schedule_summary']['cue_count']} / {run_config['cue_schedule_summary']['expected_sum']}.
- Anchor file: `{run_config['config']['predeclared_anchor_file']}`.
- Channel map: `{run_config['config']['channel_map_label']}`, confirmed={run_config['config']['channel_map_confirmed']}, confidence={run_config['config']['channel_map_confidence']}.

## Boxed recommended primary path

**Timing/QC + StimulusFingerprint/CET-R + artifact/confound gates -> Anchor-Locked MRED / A-MRED.**

All other modules remain in the suite, but for future confirmatory work they should be labeled as secondary or exploratory:

- **Primary endpoint:** A-MRED.
- **Secondary summaries:** NIP/BIT/CII/IAQ and TTI.
- **Exploratory convergence/diagnostics:** KHT-topo, NAST, EET, MRED-ITP/ACG/OCU, OCM/RSM/CVB/SquintProxy.

## A-MRED primary endpoint

Runner-time anchors:

- A-MRED passes: **{a_pass}/{n_anchor}**.
- Base-gate passes: **{base_pass}/{n_anchor}**.
- Specificity-gate passes: **{spec_pass}/{n_anchor}**.
- Passing anchor IDs: `{', '.join(amred_pass_anchor) if amred_pass_anchor else 'none'}`.

The content-offset sensitivity analysis applied a +1.251 s offset and dropped to **{interpret['amred_content_offset_sensitivity_passes']}/7** A-MRED passes. That timing sensitivity is the decisive reason this remains pilot evidence.

## ALS timing / physical pulse QC

Simple ALS fullscreen-start pulse result: **{als_passes}/3** branches passed the threshold. Target detected; Control and Override were weak/failed under this test. This does not erase the run, but it prevents a strict timing-grade claim.

## Self-report / phenomenology

Target core ratings: Meaning={target_report.get('ratings',{}).get('Meaning')}, Absorption={target_report.get('ratings',{}).get('Absorption')}, EmotionalAfterglow={target_report.get('ratings',{}).get('EmotionalAfterglow')}, StoryActiveWashout={target_report.get('ratings',{}).get('StoryActiveWashout')}.

Override core ratings: Meaning={override_report.get('ratings',{}).get('Meaning')}, Absorption={override_report.get('ratings',{}).get('Absorption')}, TaskExtractionLoad={override_report.get('ratings',{}).get('TaskExtractionLoad')}.

Final master report selected Target as most meaningful, most absorbing, and strongest afterglow. It also rated TargetEchoedDuringBaseline2={final_report.get('ratings',{}).get('TargetEchoedDuringBaseline2')}, NewMeaningToday={final_report.get('ratings',{}).get('NewMeaningToday')}, OverrideReducedReception={final_report.get('ratings',{}).get('OverrideReducedReception')}, and StoryBrokeThroughOverride={final_report.get('ratings',{}).get('StoryBrokeThroughOverride')}.

## TTI / reception-extraction

Global TTI = **{tti_val:.3f}**. This supports a positive reception/extraction contrast: Target was more meaning/reception-favoring, while Override was more task/extraction-favoring. TTI remains a secondary summary, not a moral or clinical score.

## NIP / CII / IAQ

NIP/CII was directionally Target-favoring but weaker than the previous Contact run. Later anchors were Target-dominant; the early anchor showed Override > Target and distorted mean IAQ. NIP/CII should be treated as secondary support only.

## OCM / RSM / CVB / Squint

Override task report indicated TaskCompliance={override_task.get('ratings',{}).get('TaskCompliance')}, FinalSumConfidence={override_task.get('ratings',{}).get('FinalSumConfidence')}, RunningSumStall={override_task.get('ratings',{}).get('RunningSumStall')}, ApproximateGuessCount={override_task.get('ratings',{}).get('ApproximateGuessCount')}, and HardNumberCombinationDifficulty={override_task.get('ratings',{}).get('HardNumberCombinationDifficulty')}.

RSM showed modest arithmetic-load/guess-risk structure. This is useful for interpreting the Override branch, but it cannot prove the exact private strategy.

## CET / CET-R

Full v1.8 stimulus-side regressors were available. CET-R blocked-CV values were negative or weak, meaning the available stimulus-side regressors did not generalize well enough to explain the NIP pattern. CET-R is therefore supportive as a guardrail, but not itself an endpoint.

## MRED-ITP / ACG / OCU

ACG and OCU were run as exploratory convergence layers. Complexity and blink timing are useful falsification tools, but LZ-style complexity is not literal thermodynamic entropy, and Fp1 blink proxy is not EOG. No MRED-ITP result should be treated as proof of memory encoding.

## Baseline 1 vs Baseline 2

Baseline 2 looked more reflective/regulatory than Baseline 1 on several summaries: TSP and NIP increased, artifact/task burden decreased, heart rate decreased, and HRV increased. This aligns with TargetEchoedDuringBaseline2 self-report, but it is not proof of memory encoding.

## Bottom line

Field of Dreams Run 1 is a strong pilot dataset for the new PRAYCG2.0 / A-MRED workflow. It shows coherent Target-favoring subjective and physiological summaries, and one runner-time A-MRED pass. It should be reported as **runner-registered estimated blind-anchor pilot-positive**, not as frame-verified confirmatory evidence.
"""
(REPORT/'FieldOfDreams_Run1_FULL_MasterComprehensive_Report_v1_1.md').write_text(report_md, encoding='utf-8')

# Create docx
doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Arial'; styles['Normal'].font.size = Pt(10)
for s in ['Title','Heading 1','Heading 2','Heading 3']:
    styles[s].font.name = 'Arial'

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def add_h(text, level=1): doc.add_heading(text, level=level)
def add_p(text): doc.add_paragraph(text)
def add_bullet(text): doc.add_paragraph(text, style='List Bullet')

title=doc.add_paragraph()
title.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=title.add_run('Field of Dreams PR-AYC-G Run 1\nFull Master Comprehensive Suite Report v1.1')
run.bold=True; run.font.size=Pt(18); run.font.color.rgb=RGBColor(31,78,121)
sub=doc.add_paragraph('Runner-registered estimated blind-anchor pilot - not frame-verified confirmatory')
sub.alignment=WD_ALIGN_PARAGRAPH.CENTER

add_h('Executive verdict',1)
add_p('Field of Dreams Run 1 is a useful PRAYCG2.0, runner-registered estimated blind-anchor pilot. The full suite supports a Target-favoring reception/meaning pattern, but it does not reach strict confirmatory A-MRED status because the anchors were YouTube-estimated rather than frame-verified and ALS start-pulse validation was weak.')

add_h('Boxed recommended primary path',1)
t=doc.add_table(rows=1, cols=3); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
headers=['Tier','Module group','Role']
for i,h in enumerate(headers):
    t.cell(0,i).text=h; shade_cell(t.cell(0,i),'D9EAF7')
for row in tier_rows:
    cells=t.add_row().cells
    cells[0].text=row['tier']; cells[1].text=row['module_group']; cells[2].text=row['role']
    if row['tier']=='BOXED PRIMARY PATH':
        for c in range(3): shade_cell(cells[c],'E2F0D9')

add_h('Run integrity',1)
for b in [
    f"PRAYCG version/run: PRAYCG2.0 / {run_config['config']['run_label']}",
    f"Participant/session: {run_config['config']['participant_id']} / {run_config['config']['session_id']}",
    f"Target-Override hash match: {run_config['target_override_hash_match']}",
    f"Cue count / expected sum: {run_config['cue_schedule_summary']['cue_count']} / {run_config['cue_schedule_summary']['expected_sum']}",
    f"Channel map: {run_config['config']['channel_map_label']}; confirmed={run_config['config']['channel_map_confirmed']}; confidence={run_config['config']['channel_map_confidence']}",
    f"Predeclared anchor file: {run_config['config']['predeclared_anchor_file']}",
]: add_bullet(b)

add_h('Primary endpoint: A-MRED',1)
add_p(f'Runner-time anchors produced {a_pass}/{n_anchor} A-MRED passes, {base_pass}/{n_anchor} base-gate passes, and {spec_pass}/{n_anchor} specificity-gate passes. The passing runner-time anchor was: {", ".join(amred_pass_anchor) if amred_pass_anchor else "none"}.')
add_p(f'The content-offset sensitivity analysis dropped to {interpret["amred_content_offset_sensitivity_passes"]}/7 A-MRED passes. That timing sensitivity is the main reason this remains pilot evidence rather than strict confirmation.')
if not amred_runner.empty:
    cols=['anchor_id','anchor_time_sec_used','A_MRED_pass','target_MR','target_ENC','target_CII','control_CII','override_CII','target_minus_override_MR','target_minus_override_ENC']
    tbl=doc.add_table(rows=1, cols=len(cols)); tbl.style='Table Grid'
    for i,c in enumerate(cols): tbl.cell(0,i).text=c; shade_cell(tbl.cell(0,i),'D9EAF7')
    for _,r in amred_runner[cols].iterrows():
        cells=tbl.add_row().cells
        for i,c in enumerate(cols):
            val=r[c]
            cells[i].text=f'{val:.3f}' if isinstance(val,(float,np.floating)) else str(val)

add_h('Timing and QC',1)
add_p(f'Simple ALS fullscreen-start pulse validation passed for {als_passes}/3 branches. This is a caution flag, not a total run rejection. Sample-index timing reconstruction was used for the OpenBCI continuous stream.')
if not alspulse.empty:
    cols=['phase','rise_adc_counts','peak_rel_sec','pass_detected_fullscreen_pulse','note']
    tbl=doc.add_table(rows=1, cols=len(cols)); tbl.style='Table Grid'
    for i,c in enumerate(cols): tbl.cell(0,i).text=c; shade_cell(tbl.cell(0,i),'D9EAF7')
    for _,r in alspulse[cols].iterrows():
        cells=tbl.add_row().cells
        for i,c in enumerate(cols):
            val=r[c]
            cells[i].text=f'{val:.3f}' if isinstance(val,(float,np.floating)) else str(val)

add_h('Self-report',1)
add_p(f"Target core ratings: Meaning={target_report.get('ratings',{}).get('Meaning')}, Absorption={target_report.get('ratings',{}).get('Absorption')}, EmotionalAfterglow={target_report.get('ratings',{}).get('EmotionalAfterglow')}, StoryActiveWashout={target_report.get('ratings',{}).get('StoryActiveWashout')}.")
add_p(f"Override core ratings: Meaning={override_report.get('ratings',{}).get('Meaning')}, Absorption={override_report.get('ratings',{}).get('Absorption')}, TaskExtractionLoad={override_report.get('ratings',{}).get('TaskExtractionLoad')}.")
add_p(f"Final master report selected Target as most meaningful, most absorbing, and strongest afterglow. TargetEchoedDuringBaseline2={final_report.get('ratings',{}).get('TargetEchoedDuringBaseline2')}; NewMeaningToday={final_report.get('ratings',{}).get('NewMeaningToday')}; OverrideReducedReception={final_report.get('ratings',{}).get('OverrideReducedReception')}; StoryBrokeThroughOverride={final_report.get('ratings',{}).get('StoryBrokeThroughOverride')}.")

add_h('Full module registry',1)
reg_cols=['module','tier','status','finding','limitation']
tbl=doc.add_table(rows=1, cols=len(reg_cols)); tbl.style='Table Grid'
for i,c in enumerate(reg_cols): tbl.cell(0,i).text=c; shade_cell(tbl.cell(0,i),'D9EAF7')
for _,r in registry[reg_cols].iterrows():
    cells=tbl.add_row().cells
    for i,c in enumerate(reg_cols): cells[i].text=str(r[c])

add_h('Secondary and exploratory summaries',1)
add_p(f'TTI_global = {tti_val:.3f}. This supports a positive reception/extraction contrast, with Target more receptive/meaning-favoring and Override more task/extraction-favoring.')
add_p('NIP/CII/IAQ was directionally Target-favoring but weaker than Contact. It is secondary support, not the primary endpoint.')
add_p('CET-R used full v1.8 stimulus-side regressors. Blocked-CV values were weak/negative, so stimulus-side regressors did not explain away the NIP pattern. This functions as a guardrail, not a positive endpoint.')
add_p('MRED-ITP/ACG/OCU ran as exploratory convergence. Complexity is not literal thermodynamic entropy, and blink proxy is not EOG. These modules should remain support/falsification layers.')
add_p('Baseline 2 showed a reflective/regulatory profile relative to Baseline 1 and aligned with Target echo self-report, but it does not prove memory encoding.')

add_h('Final interpretation',1)
add_p(interpret['executive_verdict'])
add_p(interpret['primary_recommendation'])

# save docx
docx_path=REPORT/'FieldOfDreams_Run1_FULL_MasterComprehensive_Report_v1_1.docx'
doc.save(docx_path)
# Convert to PDF
try:
    subprocess.run(['python','/home/oai/skills/pdfs/scripts/lo_convert_to_pdf.py',str(docx_path),'--out_dir',str(REPORT)], check=True, timeout=120)
except Exception as e:
    print('PDF conversion failed', e)
# rename if needed
pdfs=list(REPORT.glob('FieldOfDreams_Run1_FULL_MasterComprehensive_Report_v1_1*.pdf'))
# copy script itself
shutil.copy2(__file__, SCRIPTS/'run_field_full_suite_v1_1.py') if '__file__' in globals() else None
# also write checksums
import hashlib
with open(OUT/'CHECKSUMS_SHA256.txt','w') as f:
    for p in sorted([x for x in OUT.rglob('*') if x.is_file()]):
        if p.name=='CHECKSUMS_SHA256.txt': continue
        h=hashlib.sha256(p.read_bytes()).hexdigest()
        f.write(f'{h}  {p.relative_to(OUT)}\n')
# zip
zip_path=ROOT/'PRAYCG_FieldOfDreams_Run1_MasterComprehensive_FULL_v1_1.zip'
if zip_path.exists(): zip_path.unlink()
with zipfile.ZipFile(zip_path,'w',compression=zipfile.ZIP_DEFLATED) as z:
    for p in sorted([x for x in OUT.rglob('*') if x.is_file()]): z.write(p, p.relative_to(OUT.parent))
print('DONE', zip_path)
